from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)

BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "F2F4F7"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    r = sub.add_run(subtitle)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_footer(doc: Document, label: str) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = label
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor.from_string("666666")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell)
        if widths:
            set_cell_width(cell, widths[i])
        for p in cell.paragraphs:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor.from_string("0B2545")
            p.runs[0].font.size = Pt(10)
            p.paragraph_format.space_after = Pt(0)

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cells[i])
            if widths:
                set_cell_width(cells[i], widths[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def new_doc(title: str, subtitle: str, footer: str) -> Document:
    doc = Document()
    style_doc(doc)
    add_title(doc, title, subtitle)
    add_footer(doc, footer)
    return doc


def save_doc(doc: Document, filename: str) -> Path:
    path = OUT / filename
    doc.save(path)
    return path


def build_vision_scope() -> Path:
    doc = new_doc(
        "Vision and Scope",
        "Dự án: Horse Racing Tournament Website | Phiên bản 1.0 | Ngày 04/06/2026",
        "Horse Racing Tournament Website - Vision and Scope",
    )
    doc.add_heading("1. Tầm nhìn sản phẩm", level=1)
    doc.add_paragraph(
        "Horse Racing Tournament Website là hệ thống quản lý giải đua ngựa trực tuyến, "
        "hỗ trợ ban tổ chức tạo giải đấu, chủ ngựa đăng ký ngựa, nài ngựa tham gia giải, "
        "trọng tài ghi nhận kết quả và khán giả theo dõi lịch thi đấu, kết quả, bảng xếp hạng."
    )
    doc.add_paragraph(
        "Mục tiêu là giảm thao tác thủ công trong quá trình tổ chức giải, minh bạch hóa "
        "quy trình duyệt hồ sơ và công bố kết quả, đồng thời cung cấp trải nghiệm theo dõi "
        "giải đấu rõ ràng cho nhiều nhóm người dùng."
    )

    doc.add_heading("2. Vấn đề cần giải quyết", level=1)
    add_bullets(
        doc,
        [
            "Dữ liệu giải đấu, ngựa, nài ngựa, trọng tài và kết quả dễ bị phân tán nếu quản lý bằng bảng tính rời.",
            "Quy trình duyệt hồ sơ và ghép ngựa với nài ngựa cần nhiều bước xác nhận giữa Owner, Jockey và Admin.",
            "Kết quả cuộc đua cần được trọng tài nhập và Admin xác nhận trước khi dùng để tính xếp hạng.",
            "Khán giả cần xem thông tin công khai mà không có quyền chỉnh sửa dữ liệu vận hành.",
        ],
    )

    doc.add_heading("3. Phạm vi trong phiên bản này", level=1)
    add_table(
        doc,
        ["Nhóm chức năng", "Trong phạm vi"],
        [
            ["Tài khoản & vai trò", "Đăng nhập bằng mật khẩu; phân quyền Admin, Owner, Jockey, Referee, Spectator."],
            ["Quản lý giải đấu", "Admin tạo giải, tạo race, mở/đóng đăng ký, phân công trọng tài, công bố kết quả."],
            ["Quản lý ngựa", "Owner tạo và cập nhật hồ sơ ngựa, tải chứng nhận thú y, gửi đăng ký tham gia race."],
            ["Quản lý nài ngựa", "Jockey công bố hồ sơ, đăng ký tham gia tournament, nhận và phản hồi lời mời ghép cặp."],
            ["Thi đấu", "Referee kiểm tra trước race, bắt đầu race, ghi vị trí, thời gian, ghi chú và vi phạm."],
            ["Kết quả & xếp hạng", "Admin xác nhận kết quả; hệ thống tính điểm, bảng xếp hạng và trạng thái trao giải."],
            ["Thông báo", "Hệ thống gửi thông báo khi có yêu cầu duyệt, lời mời, trạng thái đăng ký hoặc kết quả."],
        ],
        [1.7, 4.8],
    )

    doc.add_heading("4. Ngoài phạm vi", level=1)
    add_bullets(
        doc,
        [
            "Không xử lý thanh toán vé hoặc bán vé trực tuyến trong phiên bản này.",
            "Không tích hợp thiết bị đo thời gian thực tế ngoài đường đua.",
            "Không tự động xác minh giấy chứng nhận thú y bằng bên thứ ba.",
            "Không hỗ trợ cá cược hoặc tính thưởng dựa trên đặt cược tiền thật.",
        ],
    )

    doc.add_heading("5. Đối tượng sử dụng", level=1)
    add_table(
        doc,
        ["Actor", "Mục tiêu chính", "Quyền chính"],
        [
            ["Admin", "Điều hành giải đấu", "Tạo giải/race, duyệt hồ sơ, phân công trọng tài, xác nhận kết quả, công bố trao giải."],
            ["Owner", "Quản lý và đăng ký ngựa", "Tạo hồ sơ ngựa, chọn Jockey đã được duyệt, đăng ký race, theo dõi kết quả."],
            ["Jockey", "Tham gia giải và điều khiển ngựa", "Công bố hồ sơ, đăng ký tournament, nhận lời mời, xem race được giao."],
            ["Referee", "Điều hành race được phân công", "Kiểm tra người tham gia, bắt đầu race, ghi nhận và nộp kết quả."],
            ["Spectator", "Theo dõi thông tin công khai", "Xem tournament, race card, kết quả, ranking và awards."],
        ],
        [1.1, 2.2, 3.2],
    )

    doc.add_heading("6. Ràng buộc và giả định", level=1)
    add_bullets(
        doc,
        [
            "Một tournament chuẩn có 10 ngựa và 10 race; tất cả 10 ngựa tham gia đủ 10 race.",
            "Chỉ kết quả đã được Admin xác nhận mới được dùng để tính bảng xếp hạng.",
            "Owner chỉ được chọn Jockey đã được Admin duyệt trong cùng tournament.",
            "Sau khi tournament hoàn tất, các vai trò chỉ có quyền xem chi tiết tournament và race.",
        ],
    )

    doc.add_heading("7. Tiêu chí thành công", level=1)
    add_bullets(
        doc,
        [
            "Người dùng thao tác được theo đúng luồng vai trò mà không thấy chức năng vượt quyền.",
            "Admin có thể hoàn tất một giải từ tạo tournament đến công bố awards.",
            "Bảng xếp hạng phản ánh đúng quy tắc điểm và tie-break.",
            "Spectator xem được dữ liệu công khai mà không cần quyền chỉnh sửa.",
        ],
    )
    return save_doc(doc, "Vision_and_Scope_Horse_Racing_Tournament_Website.docx")


def build_user_stories() -> Path:
    doc = new_doc(
        "User Stories",
        "Dự án: Horse Racing Tournament Website | Phiên bản 1.0",
        "Horse Racing Tournament Website - User Stories",
    )
    doc.add_heading("1. Epic và user stories", level=1)
    rows = [
        ["US-01", "Admin", "Là Admin, tôi muốn tạo tournament với thời gian đăng ký, ngày bắt đầu, địa điểm và quỹ giải thưởng để công bố một giải mới.", "High", "Tournament được lưu và hiển thị ở danh sách công khai."],
        ["US-02", "Admin", "Là Admin, tôi muốn tạo race trong tournament để lập lịch các vòng thi đấu.", "High", "Race có ngày, giờ, địa điểm, cự ly, surface, class và trạng thái draft/published."],
        ["US-03", "Admin", "Là Admin, tôi muốn duyệt hồ sơ ngựa để chỉ ngựa hợp lệ được đăng ký race.", "High", "Hồ sơ chuyển sang approved hoặc rejected kèm thông báo."],
        ["US-04", "Admin", "Là Admin, tôi muốn duyệt Jockey tham gia tournament để Owner chỉ chọn được Jockey hợp lệ.", "High", "Jockey được approved trong tournament tương ứng."],
        ["US-05", "Admin", "Là Admin, tôi muốn phân công một hoặc nhiều Referee cho race để có người điều hành và ghi kết quả.", "Medium", "Referee thấy race trong danh sách được giao."],
        ["US-06", "Owner", "Là Owner, tôi muốn tạo hồ sơ ngựa với thông tin giống, tuổi, sức khỏe và chứng nhận thú y để gửi duyệt.", "High", "Ngựa được tạo ở trạng thái pending và có thông báo cho Admin."],
        ["US-07", "Owner", "Là Owner, tôi muốn đăng ký ngựa vào race và chọn Jockey đã duyệt để hoàn tất entry.", "High", "Entry lưu horse, race, jockey, trạng thái xác nhận Owner/Jockey."],
        ["US-08", "Owner", "Là Owner, tôi muốn theo dõi trạng thái race entry để biết hồ sơ đang chờ Jockey hay Admin.", "Medium", "Trang Owner hiển thị pending/accepted/rejected/approved."],
        ["US-09", "Jockey", "Là Jockey, tôi muốn công bố hồ sơ cá nhân và chứng chỉ để được xét duyệt tham gia.", "High", "Profile chuyển trạng thái pending/published/rejected."],
        ["US-10", "Jockey", "Là Jockey, tôi muốn đăng ký tham gia tournament để được Owner mời lái ngựa.", "High", "Registration là unique theo tournament và jockey."],
        ["US-11", "Jockey", "Là Jockey, tôi muốn chấp nhận hoặc từ chối lời mời của Owner để xác nhận cặp Horse-Jockey.", "High", "Invitation lưu respondedAt và trạng thái accepted/rejected."],
        ["US-12", "Referee", "Là Referee, tôi muốn kiểm tra từng người tham gia là Ready hoặc Absent trước khi bắt đầu race.", "High", "Race chỉ start khi mọi participant được kiểm tra và có ít nhất một Ready."],
        ["US-13", "Referee", "Là Referee, tôi muốn nhập vị trí, finish time, ghi chú và vi phạm để nộp kết quả chính thức.", "High", "Kết quả chuyển sang submitted chờ Admin xác nhận."],
        ["US-14", "Admin", "Là Admin, tôi muốn xác nhận hoặc từ chối kết quả do Referee nộp để đảm bảo dữ liệu chính thức.", "High", "Chỉ confirmed result được dùng để ranking."],
        ["US-15", "Spectator", "Là Spectator, tôi muốn xem lịch race, race card, live status, kết quả và ranking để theo dõi giải.", "Medium", "Không có nút thao tác chỉnh sửa dữ liệu."],
        ["US-16", "System", "Là hệ thống, tôi muốn gửi notification khi có thay đổi trạng thái quan trọng để người dùng biết việc cần xử lý.", "Medium", "Notification được tạo theo userId, title, message, isRead."],
    ]
    add_table(doc, ["ID", "Actor", "User story", "Priority", "Acceptance criteria"], rows, [0.65, 0.8, 2.75, 0.75, 1.55])

    doc.add_heading("2. Definition of Done", level=1)
    add_bullets(
        doc,
        [
            "Luồng chính được triển khai ở frontend và backend hoặc có mock rõ ràng nếu là prototype.",
            "API kiểm tra role trước khi trả về hoặc thay đổi dữ liệu.",
            "Trạng thái nghiệp vụ được cập nhật nhất quán và có notification khi cần.",
            "Các màn hình chính hiển thị trạng thái rỗng, pending, approved, rejected và completed.",
        ],
    )
    return save_doc(doc, "User_Stories_Horse_Racing_Tournament_Website.docx")


def build_business_rules() -> Path:
    doc = new_doc(
        "Business Rules",
        "Dự án: Horse Racing Tournament Website | Phiên bản 1.0",
        "Horse Racing Tournament Website - Business Rules",
    )
    doc.add_heading("1. Quy tắc vai trò và truy cập", level=1)
    add_table(
        doc,
        ["ID", "Quy tắc", "Áp dụng cho", "Mức độ"],
        [
            ["BR-01", "Tài khoản chỉ được có một role chính: admin, owner, jockey, referee hoặc spectator.", "All users", "Bắt buộc"],
            ["BR-02", "Người đăng ký công khai không được tự chọn role Admin.", "Registration", "Bắt buộc"],
            ["BR-03", "Spectator chỉ được xem dữ liệu công khai và không được tạo, duyệt hoặc sửa dữ liệu.", "Spectator", "Bắt buộc"],
            ["BR-04", "Referee chỉ được thao tác trên race được Admin phân công.", "Referee", "Bắt buộc"],
        ],
        [0.7, 3.15, 1.35, 1.0],
    )

    doc.add_heading("2. Quy tắc tournament và race", level=1)
    add_table(
        doc,
        ["ID", "Quy tắc", "Ghi chú"],
        [
            ["BR-05", "Tournament tiêu chuẩn có đúng 10 ngựa và 10 race.", "Dùng cho chế độ ranking chính thức."],
            ["BR-06", "Tất cả 10 ngựa trong tournament phải tham gia đủ 10 race.", "Nếu không đủ, ranking phải thể hiện race count."],
            ["BR-07", "Race phải có ngày, giờ, địa điểm và trạng thái trước khi publish.", "Trạng thái tối thiểu: draft, registration, published, running, completed."],
            ["BR-08", "Race registration có thể đóng thủ công hoặc theo thời gian đăng ký.", "Sau khi đóng, Owner/Jockey không được thêm entry mới."],
            ["BR-09", "Gate, handicap và rating chỉ hiển thị cho Jockey sau khi race được publish.", "Tránh lộ thông tin sớm."],
        ],
        [0.7, 4.0, 1.8],
    )

    doc.add_heading("3. Quy tắc đăng ký ngựa và ghép Jockey", level=1)
    add_bullets(
        doc,
        [
            "Owner được quản lý tối đa 5 ngựa trong tài khoản.",
            "Ngựa phải được Admin duyệt trước khi đăng ký race.",
            "Jockey phải có profile hợp lệ và được duyệt trong tournament trước khi Owner có thể chọn.",
            "Một race entry phải gắn với đúng một race, một horse và một jockey.",
            "Horse-Jockey pairing cần Owner gửi lời mời, Jockey phản hồi, sau đó Admin duyệt pairing cho race nếu cần.",
        ],
    )

    doc.add_heading("4. Quy tắc trước và trong race", level=1)
    add_numbered(
        doc,
        [
            "Referee kiểm tra từng participant ở trạng thái Ready hoặc Absent.",
            "Race chỉ được bắt đầu khi race đã Published, có ít nhất một participant Ready và mọi participant đã được kiểm tra.",
            "Participant có thể bị đánh dấu disqualified nếu vi phạm; ghi chú vi phạm phải được lưu.",
            "Finish time và position phải được nhập cho các participant hoàn thành race.",
        ],
    )

    doc.add_heading("5. Quy tắc điểm và xếp hạng", level=1)
    add_table(
        doc,
        ["Vị trí", "Điểm"],
        [["1", "10"], ["2", "7"], ["3", "5"], ["4", "3"], ["5", "2"], ["6", "1"], ["7-10", "0"]],
        [1.0, 1.0],
    )
    add_numbered(
        doc,
        [
            "Ranking chỉ tính từ kết quả đã được Admin xác nhận chính thức.",
            "Sắp xếp theo tổng điểm giảm dần.",
            "Nếu bằng điểm, ưu tiên số lần về nhất nhiều hơn.",
            "Nếu vẫn bằng, ưu tiên số lần về nhì nhiều hơn.",
            "Nếu vẫn bằng, ưu tiên tổng finish time thấp hơn.",
            "Champion là ngựa xếp hạng 1 sau race thứ 10.",
        ],
    )

    doc.add_heading("6. Quy tắc hoàn tất giải", level=1)
    add_bullets(
        doc,
        [
            "Khi tournament hoàn thành, tất cả role chỉ được xem chi tiết tournament và race.",
            "Awards chỉ được publish sau khi race cuối cùng đã confirmed result.",
            "Kết quả và ranking đã publish phải được giữ ổn định để phục vụ tra cứu.",
        ],
    )
    return save_doc(doc, "Business_Rules_Horse_Racing_Tournament_Website.docx")


def build_srs_template() -> Path:
    doc = new_doc(
        "Software Requirements Specification Template",
        "Dự án: Horse Racing Tournament Website | Template có gợi ý nội dung theo IEEE-style SRS",
        "Horse Racing Tournament Website - SRS Template",
    )
    doc.add_heading("1. Introduction", level=1)
    add_table(
        doc,
        ["Mục", "Nội dung cần điền"],
        [
            ["1.1 Purpose", "Mô tả mục đích của tài liệu SRS và đối tượng đọc: nhóm phát triển, giảng viên, tester, stakeholder."],
            ["1.2 Scope", "Tóm tắt hệ thống website quản lý giải đua ngựa, các role chính và phạm vi phiên bản."],
            ["1.3 Definitions", "Giải thích Tournament, Race, Owner, Jockey, Referee, Race Entry, Ranking, Award."],
            ["1.4 References", "Liệt kê Vision and Scope, Use Case Diagram, User Stories, Business Rules, ERD, source code."],
            ["1.5 Overview", "Mô tả cấu trúc các phần còn lại của tài liệu."],
        ],
        [1.25, 5.1],
    )

    doc.add_heading("2. Overall Description", level=1)
    add_bullets(
        doc,
        [
            "Product perspective: web app gồm React frontend, Node.js backend và PostgreSQL database.",
            "Product functions: quản lý tournament/race, hồ sơ ngựa, hồ sơ Jockey, đăng ký, duyệt, live race, kết quả, ranking, notification.",
            "User classes: Admin, Owner, Jockey, Referee, Spectator.",
            "Operating environment: trình duyệt hiện đại, backend Node.js, PostgreSQL.",
            "Design constraints: phân quyền theo role, dữ liệu chính thức chỉ sau xác nhận Admin, đăng nhập dùng email/password.",
            "Assumptions and dependencies: lịch thi đấu và chứng nhận thú y được nhập thủ công; không tích hợp thiết bị đo thời gian thực tế.",
        ],
    )

    doc.add_heading("3. External Interface Requirements", level=1)
    add_table(
        doc,
        ["Interface", "Yêu cầu mô tả"],
        [
            ["User interface", "Màn hình Landing, Login/Register, Dashboard, Tournament, Race Details, Admin Panel, Owner/Jockey/Referee portal, Rankings, Results."],
            ["Software interface", "REST API giữa frontend và backend; PostgreSQL lưu users, tournaments, races, horses, invitations, entries, notifications, sessions."],
            ["Authentication interface", "Đăng nhập bằng email/password, lưu session token và kiểm tra quyền theo role."],
            ["Notification interface", "Thông báo trong hệ thống khi hồ sơ, lời mời hoặc kết quả đổi trạng thái."],
        ],
        [1.35, 5.0],
    )

    doc.add_heading("4. System Features", level=1)
    features = [
        ["SF-01", "Authentication and Authorization", "Đăng ký, đăng nhập, session, phân quyền theo role."],
        ["SF-02", "Tournament and Race Management", "Admin tạo tournament/race, publish, close registration, assign referee."],
        ["SF-03", "Horse Management", "Owner tạo hồ sơ, cập nhật thông tin, gửi duyệt và đăng ký race."],
        ["SF-04", "Jockey Management", "Jockey tạo profile, đăng ký tournament, phản hồi invitation."],
        ["SF-05", "Race Operations", "Referee pre-race check, start race, record results and violations."],
        ["SF-06", "Results and Rankings", "Admin confirm results; system calculates scores, tie-breaks, awards."],
        ["SF-07", "Public Viewing", "Spectator xem tournament, race card, results, rankings, awards."],
    ]
    add_table(doc, ["ID", "Feature", "Mô tả"], features, [0.7, 2.0, 3.65])

    doc.add_heading("5. Functional Requirements", level=1)
    add_table(
        doc,
        ["Requirement ID", "Statement", "Priority", "Source", "Acceptance Criteria"],
        [["FR-XX", "The system shall ...", "High/Medium/Low", "Use story / Business rule", "Given/When/Then hoặc điều kiện kiểm thử."]],
        [1.1, 2.6, 0.85, 1.2, 1.7],
    )

    doc.add_heading("6. Nonfunctional Requirements", level=1)
    add_table(
        doc,
        ["Category", "Requirement prompt"],
        [
            ["Security", "Mật khẩu demo cần được thay bằng hashing trong môi trường thật; API phải kiểm tra session và role."],
            ["Usability", "Màn hình theo role phải rõ trạng thái pending/approved/rejected/completed."],
            ["Reliability", "Kết quả confirmed phải không bị ghi đè ngoài quy trình Admin."],
            ["Maintainability", "Backend routes nên tách theo role; database migration lưu thay đổi schema."],
            ["Performance", "Danh sách tournament, race, ranking phản hồi trong thời gian phù hợp với lớp học/demo."],
        ],
        [1.35, 5.0],
    )

    doc.add_heading("7. Business Rules Traceability", level=1)
    add_table(
        doc,
        ["Rule ID", "Related FR", "Kiểm thử"],
        [["BR-XX", "FR-XX", "Mô tả test case xác nhận rule được áp dụng."]],
        [1.0, 1.4, 3.95],
    )

    doc.add_heading("8. Appendices", level=1)
    add_bullets(
        doc,
        [
            "Appendix A: Use Case Diagram.",
            "Appendix B: Requirements Backlog.",
            "Appendix C: ERD hoặc database schema.",
            "Appendix D: Open issues và future enhancements.",
        ],
    )
    return save_doc(doc, "SRS_Template_Horse_Racing_Tournament_Website.docx")


def build_requirements_workbook() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Unordered Requirements"
    headers = ["Req ID", "Type", "Area", "Requirement - not ordered", "Actor/Owner", "Priority", "Source", "Acceptance Hint", "Status"]
    ws.append(headers)

    rows = [
        ["NFR-03", "Nonfunctional", "Maintainability", "Backend routes should remain separated by role to keep authorization logic auditable.", "Dev Team", "Medium", "Architecture", "Route files map to Admin/Owner/Jockey/Referee/Public.", "Proposed"],
        ["FR-11", "Functional", "Referee", "The system shall let Referee mark each race participant as Ready or Absent before starting a race.", "Referee", "High", "User Stories", "Race cannot start until all participants are checked.", "Draft"],
        ["FR-03", "Functional", "Horse", "The system shall let Owner create and update horse profiles with breed, age, sex, color, weight, height, ratings, health status, notes, and veterinary certificate URL.", "Owner", "High", "Vision", "Horse is saved under ownerUserId.", "Draft"],
        ["BR-09", "Business Rule", "Ranking", "Only official confirmed race results shall be used for tournament ranking.", "System/Admin", "High", "Business Rules", "Submitted or rejected results do not affect ranking.", "Draft"],
        ["FR-17", "Functional", "Notifications", "The system shall create notifications for approval requests, invitation responses, result decisions, and important status changes.", "System", "Medium", "Role Flow", "Notification includes userId, title, message, isRead, createdAt.", "Draft"],
        ["FR-07", "Functional", "Jockey", "The system shall allow Jockey to register for a tournament and prevent duplicate registration for the same tournament.", "Jockey", "High", "Schema", "Unique tournamentId+jockeyUserId is enforced.", "Draft"],
        ["FR-01", "Functional", "Auth", "The system shall support login by email/password.", "All users", "High", "README", "Session is created for an active user account.", "Draft"],
        ["BR-02", "Business Rule", "Tournament", "A standard tournament shall contain exactly 10 horses and exactly 10 races.", "Admin/System", "High", "Ranking SRS", "Completed tournament shows 10 races and 10 horses.", "Draft"],
        ["FR-14", "Functional", "Results", "The system shall allow Admin to confirm or reject results submitted by Referee.", "Admin", "High", "User Stories", "Confirmed resultStatus becomes official.", "Draft"],
        ["FR-06", "Functional", "Jockey", "The system shall allow Jockey to publish a public profile with bio, certificate, competition level, and weight.", "Jockey", "Medium", "Role Flow", "Profile status changes from draft to pending/published.", "Draft"],
        ["NFR-01", "Nonfunctional", "Security", "The system shall restrict each role to the actions allowed for that role.", "All users", "High", "Business Rules", "Protected APIs reject unauthorized role actions.", "Draft"],
        ["FR-10", "Functional", "Admin", "The system shall allow Admin to assign one or more Referees to a race.", "Admin", "Medium", "Role Flow", "Assigned referees can view the race.", "Draft"],
        ["FR-16", "Functional", "Spectator", "The system shall let Spectator view public tournaments, race schedules, race cards, live status, results, rankings, and awards.", "Spectator", "Medium", "Vision", "No modifying controls are visible to Spectator.", "Draft"],
        ["FR-04", "Functional", "Horse Approval", "The system shall allow Admin to approve or reject horse registration requests.", "Admin", "High", "Role Flow", "Owner receives status update.", "Draft"],
        ["BR-06", "Business Rule", "Race Entry", "Owner may select only Jockeys approved for the same tournament.", "Owner", "High", "ERD", "Selection list excludes unapproved jockeys.", "Draft"],
        ["FR-13", "Functional", "Race Operations", "The system shall allow Referee to record position, finish time, notes, and violation notes for race entries.", "Referee", "High", "Role Flow", "Submitted results include all required fields.", "Draft"],
        ["FR-02", "Functional", "Tournament", "The system shall allow Admin to create tournaments with name, status, registration window, dates, location, and prize pool.", "Admin", "High", "README/Schema", "Tournament appears in tournament list.", "Draft"],
        ["NFR-02", "Nonfunctional", "Usability", "The UI shall show clear states for pending, approved, rejected, published, running, completed, and awards published.", "All users", "Medium", "Role Flow", "Status badges are visible on relevant screens.", "Draft"],
        ["BR-11", "Business Rule", "Ranking", "Ranking tie-breaks shall use total score, first-place count, second-place count, then total finish time ascending.", "System", "High", "Ranking SRS", "Tie scenario sorts as specified.", "Draft"],
        ["FR-09", "Functional", "Admin", "The system shall allow Admin to close race registration and publish a race.", "Admin", "High", "Role Flow", "No new entries after close; published race visible to participants.", "Draft"],
        ["FR-05", "Functional", "Race Entry", "The system shall allow Owner to register an approved horse for a race and invite/select a Jockey.", "Owner", "High", "User Stories", "Race entry links raceId, horseId, jockeyUserId.", "Draft"],
        ["BR-01", "Business Rule", "Roles", "A user account shall have exactly one primary role: Admin, Owner, Jockey, Referee, or Spectator.", "All users", "High", "Schema", "Role value is one of the allowed set.", "Draft"],
        ["FR-15", "Functional", "Ranking", "The system shall calculate and display tournament ranking by selected tournament.", "System/Spectator", "High", "Ranking SRS", "Changing selector recalculates table.", "Draft"],
        ["FR-08", "Functional", "Invitations", "The system shall allow Jockey to accept or reject Owner invitations for a race.", "Jockey", "High", "Role Flow", "Invitation status and respondedAt update.", "Draft"],
        ["BR-14", "Business Rule", "Completion", "When a tournament is completed, all roles shall only view tournament and race details.", "All users", "High", "Ranking SRS", "Action buttons are hidden/disabled.", "Draft"],
        ["FR-12", "Functional", "Referee", "The system shall allow Referee to start a race only when it is Published, at least one participant is Ready, and all participants are checked.", "Referee", "High", "Business Rules", "Start action is blocked otherwise.", "Draft"],
    ]
    for row in rows:
        ws.append(row)

    header_fill = PatternFill("solid", fgColor="1F4D78")
    thin = Side(style="thin", color="D9E2EC")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

    widths = [12, 16, 18, 58, 18, 12, 18, 42, 12]
    for i, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:I{ws.max_row}"

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
        row[5].alignment = Alignment(horizontal="center", vertical="top", wrap_text=True)
        row[8].alignment = Alignment(horizontal="center", vertical="top", wrap_text=True)

    legend = wb.create_sheet("Legend")
    legend.append(["Field", "Meaning"])
    legend_rows = [
        ["Type", "Functional, Nonfunctional, or Business Rule."],
        ["Priority", "High = must have for SRS baseline; Medium = useful for complete workflow; Low = later enhancement."],
        ["Status", "Draft means extracted requirement still needs stakeholder confirmation."],
        ["Unordered", "Rows are intentionally not grouped by ID so the worksheet can be used for ordering/prioritization exercises."],
    ]
    for row in legend_rows:
        legend.append(row)
    for cell in legend[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center")
    legend.column_dimensions["A"].width = 18
    legend.column_dimensions["B"].width = 90
    for row in legend.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

    path = OUT / "Requirements_Backlog_Unordered_Horse_Racing_Tournament_Website.xlsx"
    wb.save(path)
    return path


def load_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            pass
    return ImageFont.load_default()


def text_center(draw: ImageDraw.ImageDraw, xy, text: str, font, fill="black", max_width=180, line_spacing=6):
    x, y, w, h = xy
    words = text.split()
    lines = []
    current = ""
    for word in words:
        probe = f"{current} {word}".strip()
        if draw.textlength(probe, font=font) <= max_width:
            current = probe
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    line_heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        line_heights.append(bbox[3] - bbox[1])
    total_h = sum(line_heights) + line_spacing * (len(lines) - 1)
    yy = y + (h - total_h) / 2
    for line, lh in zip(lines, line_heights):
        tw = draw.textlength(line, font=font)
        draw.text((x + (w - tw) / 2, yy), line, font=font, fill=fill)
        yy += lh + line_spacing


def draw_actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str, font) -> tuple[int, int]:
    draw.ellipse((x + 35, y, x + 65, y + 30), outline="#0B2545", width=3)
    draw.line((x + 50, y + 30, x + 50, y + 85), fill="#0B2545", width=3)
    draw.line((x + 15, y + 50, x + 85, y + 50), fill="#0B2545", width=3)
    draw.line((x + 50, y + 85, x + 20, y + 125), fill="#0B2545", width=3)
    draw.line((x + 50, y + 85, x + 80, y + 125), fill="#0B2545", width=3)
    text_center(draw, (x - 15, y + 130, 130, 42), label, font, "#0B2545", 130)
    return x + 50, y + 65


def draw_use_case(draw: ImageDraw.ImageDraw, center: tuple[int, int], label: str, font) -> tuple[int, int]:
    cx, cy = center
    w, h = 250, 82
    box = (cx - w // 2, cy - h // 2, cx + w // 2, cy + h // 2)
    draw.ellipse(box, fill="#FFFFFF", outline="#2E74B5", width=3)
    text_center(draw, (box[0] + 10, box[1] + 8, w - 20, h - 16), label, font, "#0B2545", w - 34)
    return center


def draw_line(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line((start[0], start[1], end[0], end[1]), fill="#6B7280", width=2)


def build_use_case_image() -> tuple[Path, Path]:
    img = Image.new("RGB", (1800, 1200), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = load_font(34, True)
    actor_font = load_font(20, True)
    uc_font = load_font(18)
    small_font = load_font(16)

    draw.text((70, 45), "Use Case Diagram - Horse Racing Tournament Website", font=title_font, fill="#0B2545")
    boundary = (380, 120, 1420, 1090)
    draw.rounded_rectangle(boundary, radius=24, outline="#1F4D78", width=4, fill="#F8FAFC")
    draw.text((410, 140), "Horse Racing Tournament Website", font=load_font(24, True), fill="#1F4D78")

    actors = {
        "Spectator": draw_actor(draw, 90, 240, "Spectator", actor_font),
        "Owner": draw_actor(draw, 90, 550, "Owner", actor_font),
        "Jockey": draw_actor(draw, 90, 835, "Jockey", actor_font),
        "Admin": draw_actor(draw, 1540, 250, "Admin", actor_font),
        "Referee": draw_actor(draw, 1540, 760, "Referee", actor_font),
    }

    use_cases = {
        "auth": draw_use_case(draw, (650, 235), "Đăng ký / Đăng nhập", uc_font),
        "public": draw_use_case(draw, (650, 370), "Xem tournament, race card, kết quả", uc_font),
        "ranking": draw_use_case(draw, (650, 505), "Xem ranking và awards", uc_font),
        "horse": draw_use_case(draw, (650, 660), "Quản lý hồ sơ ngựa", uc_font),
        "entry": draw_use_case(draw, (650, 800), "Đăng ký ngựa vào race", uc_font),
        "jprofile": draw_use_case(draw, (650, 940), "Công bố hồ sơ Jockey", uc_font),
        "join": draw_use_case(draw, (965, 940), "Đăng ký tham gia tournament", uc_font),
        "create": draw_use_case(draw, (1110, 240), "Tạo tournament và race", uc_font),
        "approve": draw_use_case(draw, (1110, 385), "Duyệt hồ sơ, đăng ký, pairing", uc_font),
        "publish": draw_use_case(draw, (1110, 530), "Đóng đăng ký và publish race", uc_font),
        "assign": draw_use_case(draw, (1110, 675), "Phân công Referee", uc_font),
        "check": draw_use_case(draw, (1110, 810), "Pre-race check", uc_font),
        "record": draw_use_case(draw, (1110, 945), "Ghi và nộp kết quả race", uc_font),
        "confirm": draw_use_case(draw, (1110, 1080), "Xác nhận kết quả chính thức", uc_font),
    }

    links = [
        ("Spectator", "auth"), ("Spectator", "public"), ("Spectator", "ranking"),
        ("Owner", "auth"), ("Owner", "horse"), ("Owner", "entry"), ("Owner", "ranking"),
        ("Jockey", "auth"), ("Jockey", "jprofile"), ("Jockey", "join"), ("Jockey", "ranking"),
        ("Admin", "create"), ("Admin", "approve"), ("Admin", "publish"), ("Admin", "assign"), ("Admin", "confirm"),
        ("Referee", "check"), ("Referee", "record"),
    ]
    for actor, uc in links:
        draw_line(draw, actors[actor], use_cases[uc])

    draw.line((965, 810, 1110, 810), fill="#6B7280", width=2)
    draw.text((820, 825), "<<include>>", font=small_font, fill="#6B7280")
    draw.line((1110, 945, 1110, 1038), fill="#6B7280", width=2)
    draw.text((1125, 1002), "<<submit>>", font=small_font, fill="#6B7280")
    draw.line((1110, 530, 1110, 344), fill="#9CA3AF", width=1)
    draw.text((1130, 455), "<<requires approval>>", font=small_font, fill="#6B7280")

    png_path = OUT / "Use_Case_Diagram_Horse_Racing_Tournament_Website.png"
    img.save(png_path)

    svg_path = OUT / "Use_Case_Diagram_Horse_Racing_Tournament_Website.svg"
    svg_path.write_text(
        """<svg xmlns="http://www.w3.org/2000/svg" width="1800" height="1200" viewBox="0 0 1800 1200">
  <rect width="1800" height="1200" fill="#ffffff"/>
  <image href="Use_Case_Diagram_Horse_Racing_Tournament_Website.png" x="0" y="0" width="1800" height="1200"/>
</svg>
""",
        encoding="utf-8",
    )
    return png_path, svg_path


def main() -> None:
    outputs = [
        build_vision_scope(),
        build_user_stories(),
        build_business_rules(),
        build_srs_template(),
        build_requirements_workbook(),
    ]
    outputs.extend(build_use_case_image())
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
